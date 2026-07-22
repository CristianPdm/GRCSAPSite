"""Generacion del Informe de Auditoria GRC-SOD en formato Word (.docx),
equivalente a expWord() de SAP_SOD_Analyzer_v640.html.

El original armaba el XML OOXML a mano (helpers wr/wp/wTbl) y lo empaquetaba
con JSZip en el navegador. Aqui se usa python-docx en el servidor, que
genera el mismo tipo de documento (titulo, tablas con color/negrita,
saltos de pagina) con una API de mas alto nivel. Los datos se obtienen de
las mismas funciones del motor SOD/Licencias que usa el Informe HTML
(informe.html) -- no se recalcula nada, solo se reordena en parrafos/tablas.
"""
import io
from datetime import date

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.sod.engine import build_audit_report_data, build_user_risk

COLOR_PRIMARY = RGBColor(0x1A, 0x56, 0xDB)
COLOR_DARK = RGBColor(0x2D, 0x37, 0x48)
COLOR_MUTED = RGBColor(0x6B, 0x72, 0x80)
COLOR_CRIT = RGBColor(0xC8, 0x1F, 0x1F)
COLOR_HIGH = RGBColor(0xC4, 0x5A, 0x00)
COLOR_MED = RGBColor(0x8A, 0x62, 0x00)
COLOR_OK = RGBColor(0x06, 0x5F, 0x46)
COLOR_ADV = RGBColor(0x6B, 0x21, 0xA8)
COLOR_CORE = RGBColor(0x1E, 0x40, 0xAF)
HEADER_BG = "EAF0FB"
ZEBRA_BG = "F7FAFC"


def _shade_cell(cell, hex_fill):
    """python-docx no trae un helper para el color de fondo de una celda;
    se inyecta el <w:shd> directamente sobre tcPr, igual que hacia wCell()
    con w:shd en el original."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(
        qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_fill}
    )
    tc_pr.append(shd)


def _set_cell(cell, text, bold=False, color=None, mono=False, size=9):
    run = cell.paragraphs[0].add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if mono:
        run.font.name = "Consolas"


def _table(doc, widths_in, headers, rows):
    """rows: lista de listas de celdas; cada celda es un str o un dict
    {text, bold, color, mono}."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        _set_cell(cell, header, bold=True, color=COLOR_DARK, size=8)
        _shade_cell(cell, HEADER_BG)
        cell.width = Inches(widths_in[i])
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, value in enumerate(row):
            obj = value if isinstance(value, dict) else {"text": value}
            _set_cell(cells[col_idx], obj.get("text", ""), bold=obj.get("bold", False),
                      color=obj.get("color"), mono=obj.get("mono", False))
            if row_idx % 2 == 1:
                _shade_cell(cells[col_idx], ZEBRA_BG)
            cells[col_idx].width = Inches(widths_in[col_idx])
    return table


def _h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = COLOR_PRIMARY


def _h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_DARK


def _note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED


def build_informe_docx():
    """Construye el .docx completo y devuelve un BytesIO listo para enviar
    como descarga (igual contenido que sod/informe.html, mas la seccion de
    Optimizacion FUE que en el HTML vive en una pantalla aparte)."""
    from app.licenses.engine import compute_fue_optimization

    data = build_audit_report_data()
    opt = compute_fue_optimization()
    s = data["summary"]
    fs = data["fue_stats"]

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = Inches(0.75)
    section.left_margin = section.right_margin = Inches(0.75)

    fecha = date.today().strftime("%d/%m/%Y")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GRUPO SIMPA S.A.")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_PRIMARY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Informe de Auditoría — Segregación de Funciones & Licencias FUE")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_DARK

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Fecha: {fecha}   |   SAP S/4HANA Private Cloud 2022   |   CONFIDENCIAL")
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED
    doc.add_paragraph()

    # 1. Resumen Ejecutivo
    _h1(doc, "1. Resumen Ejecutivo")
    _table(
        doc, [1.2, 1.2, 1.2, 1.2, 1.0],
        ["Usuarios activos", "SOD Críticos", "SOD Altos", "SOD Medios", "Total FUE"],
        [[
            {"text": data["total_usuarios_activos"], "bold": True, "color": COLOR_PRIMARY},
            {"text": s["critical"], "bold": True, "color": COLOR_CRIT if s["critical"] else COLOR_OK},
            {"text": s["high"], "bold": True, "color": COLOR_HIGH if s["high"] else COLOR_OK},
            {"text": s["medium"], "bold": True, "color": COLOR_MED if s["medium"] else COLOR_OK},
            {"text": f"{fs['total_fue']:.2f}", "bold": True, "color": COLOR_PRIMARY},
        ]],
    )
    doc.add_paragraph()
    _note(doc, f"FUE: {fs['ADV']} GB Advanced · {fs['CORE']} GC Core · {fs['SELF']} GD Self-Service")
    if s["critical"] > 0:
        concl, ccolor = f"CRITICO: Se detectaron {s['critical']} conflicto(s) SOD críticos. Acción inmediata requerida.", COLOR_CRIT
    elif s["high"] > 0:
        concl, ccolor = f"ALTO: {s['high']} conflicto(s) altos a remediar en 30 días.", COLOR_HIGH
    else:
        concl, ccolor = f"Sin conflictos críticos ni altos. {s['medium']} conflicto(s) medios requieren controles compensatorios.", COLOR_OK
    p = doc.add_paragraph()
    run = p.add_run(concl)
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = ccolor
    doc.add_page_break()

    # 2. Conflictos SOD Detectados
    _h1(doc, "2. Conflictos SOD Detectados")
    grupos = [
        ("CRITICO", data["criticos_activos"], COLOR_CRIT),
        ("ALTO", data["altos_activos"], COLOR_HIGH),
        ("MEDIO", data["medios_activos"], COLOR_MED),
    ]
    if not any(items for _, items, _ in grupos):
        _note(doc, "No se detectaron conflictos SOD en la base de datos analizada.")
    else:
        for nivel, items, color in grupos:
            if not items:
                continue
            _h2(doc, f"{nivel} — {len(items)} conflicto(s)")
            rows = []
            for item in items:
                rule = item["rule"]
                roles = data["roles_by_rule_id"].get(rule.id, [])
                t1, t2 = rule.tcodes1_list(), rule.tcodes2_list()
                conflicted = item["conflicted_users"]
                rows.append([
                    {"text": rule.id, "bold": True, "color": color, "mono": True},
                    {"text": rule.descripcion},
                    {"text": ", ".join(t1[:6]) + ("…" if len(t1) > 6 else ""), "mono": True},
                    {"text": ", ".join(t2[:6]) + ("…" if len(t2) > 6 else ""), "mono": True},
                    {"text": str(len(roles))},
                    {"text": str(len(conflicted)), "bold": bool(conflicted), "color": color if conflicted else None},
                ])
            _table(doc, [0.5, 2.1, 1.2, 1.2, 0.5, 0.6],
                   ["ID", "Descripción", "TCodes Lado A", "TCodes Lado B", "Roles", "Usuarios"], rows)
            doc.add_paragraph()
    doc.add_page_break()

    # 3. Usuarios en Riesgo
    _h1(doc, "3. Usuarios en Riesgo")
    user_risk = build_user_risk()
    top_u = [u for u in user_risk if u["cc"] > 0][:60]
    has_meta = data["fue_source_oficial"]
    if not top_u:
        _note(doc, "No se encontraron usuarios con conflictos SOD activos.")
    else:
        _note(doc, f"{len(top_u)} usuario(s) con conflictos SOD activos.")
        doc.add_paragraph()
        headers = ["Usuario"]
        widths = [1.0]
        if has_meta:
            headers.append("Nombre")
            widths.append(1.6)
        headers.append("Nivel")
        widths.append(0.7)
        headers.append("FUE SAP (oficial)" if has_meta else "FUE por roles")
        widths.append(1.1)
        if has_meta:
            headers.append("Último acceso")
            widths.append(0.9)
        headers.append("Conflictos")
        widths.append(1.6 if has_meta else 3.7)

        nivc = {"CRITICO": COLOR_CRIT, "ALTO": COLOR_HIGH, "MEDIO": COLOR_MED, "SIN": COLOR_OK}
        rows = []
        for u in top_u:
            row = [{"text": u["user"], "bold": True, "mono": True}]
            if has_meta:
                row.append({"text": u["nombre"] or "—"})
            row.append({"text": u["mx"], "bold": True, "color": nivc.get(u["mx"])})
            row.append({"text": (u["fue_from_db_label"] if has_meta else u["fue_label"]) or "—"})
            if has_meta:
                la = u["last_access"]
                la_text = f"{la.strftime('%Y-%m-%d')} ({u['days']}d)" if la else "—"
                la_color = COLOR_CRIT if u["days"] and u["days"] > 90 else (COLOR_HIGH if u["days"] and u["days"] > 30 else None)
                row.append({"text": la_text, "mono": True, "color": la_color})
            conflicts = u["conflicts"]
            row.append({"text": ", ".join(conflicts[:5]) + (f" +{len(conflicts) - 5}" if len(conflicts) > 5 else "")})
            rows.append(row)
        _table(doc, widths, headers, rows)
    doc.add_page_break()

    # 4. Analisis de Optimizacion FUE
    _h1(doc, "4. Análisis de Optimización FUE")
    _note(doc, f"Ahorro total estimado: {opt['total_savings']:.3f} FUE (inactivaciones: "
               f"{opt['total_inactive_savings']:.3f} · downgrades: {opt['total_downgrade_savings']:.3f})")
    doc.add_paragraph()
    if opt["inactive"]:
        _h2(doc, f"4.1 Candidatos a baja por inactividad — {len(opt['inactive'])} usuarios (>90 días)")
        rows = []
        for u in opt["inactive"][:50]:
            fue_lower = (u["fue_type"] or "").lower()
            fcolor = COLOR_ADV if "adv" in fue_lower else (COLOR_CORE if "core" in fue_lower else COLOR_OK)
            dcolor = COLOR_CRIT if u["days"] > 365 else (COLOR_HIGH if u["days"] > 180 else COLOR_MED)
            rows.append([
                {"text": u["user"], "bold": True, "mono": True},
                {"text": u["nombre"] or "—"},
                {"text": u["fue_type"], "color": fcolor},
                {"text": f"{u['fue_weight']:.3f}", "bold": True, "color": COLOR_CRIT},
                {"text": u["last_access"].strftime("%Y-%m-%d") if u["last_access"] else "—", "mono": True},
                {"text": f"{u['days']}d", "bold": True, "color": dcolor},
            ])
        _table(doc, [0.9, 1.3, 1.3, 0.6, 0.9, 0.6],
               ["Usuario", "Nombre", "FUE asignado", "Peso FUE", "Último acceso", "Días"], rows)
        doc.add_paragraph()
    if opt["downgrade"]:
        _h2(doc, f"4.2 Candidatos a downgrade — {len(opt['downgrade'])} usuarios")
        rows = []
        for u in opt["downgrade"][:50]:
            rows.append([
                {"text": u["user"], "bold": True, "mono": True},
                {"text": u["nombre"] or "—"},
                {"text": u["current_fue"], "color": COLOR_CRIT},
                {"text": u["derived_label"] or "—"},
                {"text": u["suggested_fue"], "color": COLOR_OK},
                {"text": f"+{u['savings']:.3f}", "bold": True, "color": COLOR_OK},
            ])
        _table(doc, [0.9, 1.3, 1.3, 0.9, 0.9, 0.6],
               ["Usuario", "Nombre", "FUE actual (SAP)", "FUE por roles", "FUE sugerido", "Ahorro"], rows)
        doc.add_paragraph()
    if not opt["inactive"] and not opt["downgrade"]:
        _note(doc, "No se identificaron oportunidades de optimización FUE. Importa FUE_Users.xlsx "
                   "desde el módulo de Licencias para habilitar este análisis.")
    doc.add_page_break()

    # 5. Plan de Remediacion
    _h1(doc, "5. Plan de Remediación")
    acciones = []
    if s["critical"] > 0:
        acciones.append([
            {"text": f"Revocar {s['critical']} conflicto(s) CRITICO", "bold": True, "color": COLOR_CRIT},
            "Resp. SAP + Gerencia", "0-15 días", "Sección 2",
        ])
    if s["high"] > 0:
        acciones.append([f"Remediar {s['high']} conflicto(s) ALTO", "Resp. SAP", "30 días", "Sección 2"])
    if opt["inactive"]:
        acciones.append([f"Dar de baja {len(opt['inactive'])} usuarios inactivos", "Resp. SAP + RRHH", "15 días", "Sección 4.1"])
    if opt["downgrade"]:
        acciones.append([f"Revisar downgrade {len(opt['downgrade'])} usuarios", "Resp. SAP + Licencias", "30 días", "Sección 4.2"])
    acciones.append(["Activar SM20/SM21 para roles críticos", "Basis / BC", "Inmediato", "Política BC"])
    acciones.append(["Documentar excepciones SOD", "Auditoría Interna", "30 días", "Política GRC"])
    acciones.append(["Re-análisis SOD trimestral", "Resp. SAP", "Trimestral", "Proceso GRC"])
    _table(doc, [2.4, 1.7, 1.0, 1.4], ["Acción", "Responsable", "Plazo", "Referencia"], acciones)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
